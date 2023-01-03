# -*- coding:UTF-8 -*-
import os.path


class MyFileUtil:
    __writeCount = 0

    def __int__(self, writeCount):
        self.__writeCount = writeCount

    def readYaml(self, path, fields=''):
        import yaml
        with open(path, 'a+', encoding="UTF-8") as file:
            file.seek(0)
            data = file.read()
            r = yaml.load(data, Loader=yaml.FullLoader)
            res = self.__paring_fields(r, fields)
        return res

    @staticmethod
    def __paring_fields(res, fields):
        if fields:
            f = fields.split('.')
            for field in f:
                res = res[field]
        return res

    @staticmethod
    def getPath(filename, path=''):
        '''获取最终文件路径'''
        import os
        filename = os.path.normpath(filename)
        filename_path = os.path.dirname(filename)
        path = os.path.join(path, filename_path)
        if path:
            path = os.path.normpath(path)
            abspath = os.path.abspath(path)
            if os.path.isfile(path):
                # raise RuntimeError('path传入了一个文件')
                return False
        else:
            abspath = os.getcwd()
        if not os.path.exists(abspath):
            os.makedirs(abspath)
        filename_name = os.path.basename(filename)
        if not filename_name:
            raise RuntimeError('没有文件名')
        fullpath = os.path.join(abspath, filename_name)
        return fullpath

    def newFile(self, filename, path=''):
        fullpath = self.getPath(filename, path)
        # if os.path.exists(fullpath):
        #     # raise RuntimeError('该路径中已存在该文件')
        #     return
        with open(fullpath, mode='a+', encoding='UTF-8') as fp:
            fp.close()
        return fullpath

    @staticmethod
    def writeFile(filepath, content=''):
        import os
        if os.path.exists(filepath):
            with open(filepath, mode='a+', encoding='UTF-8') as fp:
                fp.write(content)
                res = len(content)
                return res
        else:
            res = 0
            return res

    def writeDocx(self, path, content, title='', date=''):
        from docx import Document
        import datetime
        if os.path.exists(path):
            wordfile = Document(path)
        else:
            wordfile = Document()
        if not date:
            # 默认当日日期
            date = datetime.date.today()

        title2 = wordfile.add_heading(str(date) + str(title), level=2)
        paragraph = wordfile.add_paragraph(content)
        end_content = ''
        wordfile.add_paragraph(end_content)

        self.__writeCount += 1
        print('导入%s第%d篇新闻' % (date, self.__writeCount))

        wordfile.save(path)


if __name__ == '__main__':
    myFile = MyFileUtil()
    result = myFile.readYaml('test.yaml')
    filepath = myFile.newFile('a/')
    print(filepath)
